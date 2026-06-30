#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""把开题报告 Markdown 正文渲染为结构化中文 .docx。

设计要点：
- 中文字体设为宋体（正文）/黑体（标题），西文 Times New Roman；
- 解析 markdown 的标题(#)、表格(|)、列表、加粗(**)、普通段落；
- 表格用于封面信息与正文中的小表；
- 参考文献段落使用悬挂缩进与小一号字。
仅依赖 python-docx。
"""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
SRC = HERE / "开题报告_正文.md"
OUT = HERE / "开题报告.docx"

CJK_BODY = "宋体"
CJK_HEAD = "黑体"
LATIN = "Times New Roman"


def set_run_font(run, latin=LATIN, cjk=CJK_BODY, size=None, bold=None, color=None):
    """设置 run 的中西文字体（python-docx 需手动写 eastAsia）。"""
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cjk)


def add_runs_with_bold(par, text, **font_kw):
    """解析 **加粗** 标记，逐段添加 run。"""
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = par.add_run(p[2:-2])
            kw = dict(font_kw)
            kw["bold"] = True
            set_run_font(r, **kw)
        else:
            r = par.add_run(p)
            set_run_font(r, **font_kw)


def add_heading(doc, text, level):
    sizes = {1: 16, 2: 14, 3: 13, 4: 12}
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    par.paragraph_format.space_after = Pt(6)
    par.paragraph_format.keep_with_next = True
    r = par.add_run(text)
    set_run_font(r, latin=LATIN, cjk=CJK_HEAD, size=sizes.get(level, 12), bold=True,
                 color=RGBColor(0, 0, 0))
    return par


def add_body(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True):
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.alignment = align
    pf.line_spacing = 1.5
    pf.space_after = Pt(4)
    if first_indent:
        pf.first_line_indent = Pt(2 * size)  # 首行缩进 2 字符
    add_runs_with_bold(par, text, latin=LATIN, cjk=CJK_BODY, size=size)
    return par


def add_ref(doc, text):
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.3
    pf.space_after = Pt(2)
    pf.left_indent = Cm(0.74)
    pf.first_line_indent = Cm(-0.74)  # 悬挂缩进
    add_runs_with_bold(par, text, latin=LATIN, cjk=CJK_BODY, size=10.5)


def style_table(table):
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                par.paragraph_format.line_spacing = 1.2
                for r in par.runs:
                    set_run_font(r, latin=LATIN, cjk=CJK_BODY, size=10.5)


def flush_table(doc, rows):
    """rows: list of list[str]; 第一行作表头加粗。"""
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    for i, row in enumerate(rows):
        for j in range(ncol):
            cell = table.cell(i, j)
            txt = row[j] if j < len(row) else ""
            cell.paragraphs[0].text = ""
            add_runs_with_bold(cell.paragraphs[0], txt, latin=LATIN, cjk=CJK_BODY,
                               size=10.5)
            if i == 0:
                for r in cell.paragraphs[0].runs:
                    r.bold = True
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_separator_row(line):
    return bool(re.match(r"^\|?\s*:?-{2,}", line)) and set(line) <= set("|-: ")


def main():
    md = SRC.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()
    # 页边距
    for s in doc.sections:
        s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.0); s.right_margin = Cm(3.0)
    # 默认样式字体
    normal = doc.styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().append(OxmlElement("w:rFonts"))
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_BODY)

    i = 0
    in_refs = False
    table_buf = []

    def flush_buf():
        nonlocal table_buf
        if table_buf:
            flush_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        # 表格行累积
        if raw.startswith("|"):
            if is_separator_row(raw):
                i += 1
                continue
            table_buf.append(parse_table_row(raw))
            i += 1
            continue
        else:
            flush_buf()

        stripped = raw.strip()
        if not stripped:
            i += 1
            continue

        # 水平线
        if stripped == "---":
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if "参考文献" in text and level == 2:
                in_refs = True
            add_heading(doc, text, level)
            i += 1
            continue

        # 参考文献编号项
        if in_refs and re.match(r"^\d+\.\s", stripped):
            add_ref(doc, stripped)
            i += 1
            continue

        # 列表项（- 或 数字.）作为普通段落但带项目符号缩进
        if re.match(r"^[-*]\s+", stripped):
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.74)
            par.paragraph_format.line_spacing = 1.5
            par.paragraph_format.space_after = Pt(3)
            add_runs_with_bold(par, "• " + stripped[2:], latin=LATIN, cjk=CJK_BODY, size=12)
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped) and not in_refs:
            add_body(doc, stripped, size=12, first_indent=False)
            i += 1
            continue

        # 斜体注释行
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            par = doc.add_paragraph()
            par.paragraph_format.line_spacing = 1.3
            r = par.add_run(stripped.strip("*"))
            set_run_font(r, latin=LATIN, cjk=CJK_BODY, size=10.5)
            r.italic = True
            i += 1
            continue

        # 普通段落
        add_body(doc, stripped, size=12)
        i += 1

    flush_buf()
    doc.save(OUT)
    print(f"已生成: {OUT}")


if __name__ == "__main__":
    main()
