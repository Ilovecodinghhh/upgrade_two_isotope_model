#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Export manuscript Markdown drafts to Word documents.

This is a lightweight Markdown-to-docx exporter for the current manuscript
drafts. It preserves headings, paragraphs, simple tables, fenced code blocks,
and local PNG/JPEG figures used by the manuscript and Supplementary
Information. It is intentionally conservative and avoids changing the source
Markdown.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - exercised only on missing env
    raise SystemExit(
        "python-docx is required to export Word documents. "
        "Install python-docx or use Pandoc."
    ) from exc


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "word_exports"
IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")


@dataclass(frozen=True)
class ExportJob:
    markdown: Path
    docx: Path
    title: str


def add_inline_markdown(paragraph, text: str) -> None:
    """Add text with simple bold and inline-code styling."""
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def is_table_delimiter(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if is_table_delimiter(line):
            continue
        rows.append([cell.strip().strip("`") for cell in line.strip().strip("|").split("|")])
    return rows


def add_table(document: Document, lines: list[str]) -> None:
    rows = parse_table(lines)
    if not rows:
        return
    n_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(n_cols):
            value = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = value
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    document.add_paragraph()


def add_image(document: Document, md_path: Path, line: str) -> bool:
    match = IMAGE_RE.search(line)
    if not match:
        return False
    rel_path = match.group("path").strip()
    alt_text = match.group("alt").strip()
    image_path = (md_path.parent / rel_path).resolve()
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Missing figure: {rel_path}]").italic = True
        return True

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(6.2))
    except Exception:
        run.add_picture(str(image_path), width=Inches(5.2))
    if alt_text:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(alt_text).italic = True
    return True


def flush_paragraph(document: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer).strip()
    if text:
        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, text)
    buffer.clear()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for level in range(1, 4):
        styles[f"Heading {level}"].font.name = "Times New Roman"


def export_markdown_to_docx(md_path: Path, out_path: Path, title: str) -> None:
    text = md_path.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    if title:
        document.core_properties.title = title

    paragraph_buffer: list[str] = []
    table_buffer: list[str] = []
    in_code = False
    code_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_table(document, table_buffer)
            table_buffer = []

    def flush_code() -> None:
        nonlocal code_buffer
        if code_buffer:
            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code_buffer))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            code_buffer = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            flush_paragraph(document, paragraph_buffer)
            flush_table()
            continue

        if line.lstrip().startswith("|") and "|" in line.strip()[1:]:
            flush_paragraph(document, paragraph_buffer)
            table_buffer.append(line)
            continue

        flush_table()

        if add_image(document, md_path, line):
            flush_paragraph(document, paragraph_buffer)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush_paragraph(document, paragraph_buffer)
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2), level=level)
            continue

        if line.startswith("- ") or re.match(r"^\d+[.)]\s+", line):
            flush_paragraph(document, paragraph_buffer)
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, re.sub(r"^(-|\d+[.)])\s+", "", line))
            continue

        paragraph_buffer.append(line)

    flush_paragraph(document, paragraph_buffer)
    flush_table()
    if in_code:
        flush_code()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(out_path)


def main() -> int:
    jobs = [
        ExportJob(
            ROOT / "manuscript_first_draft.md",
            OUTPUT_DIR / "manuscript_first_draft.docx",
            "Seasonal Methane Isotope Phasors Constrain the Methane OH 13C Kinetic Isotope Effect",
        ),
        ExportJob(
            ROOT / "supplementary_information_first_draft.md",
            OUTPUT_DIR / "supplementary_information_first_draft.docx",
            "Supplementary Information",
        ),
    ]
    for job in jobs:
        if not job.markdown.exists():
            raise FileNotFoundError(job.markdown)
        export_markdown_to_docx(job.markdown, job.docx, job.title)
        print(job.docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
